from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
section.left_margin = Inches(1.1)
section.right_margin = Inches(1.1)

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BLUE  = RGBColor(0x1A, 0x56, 0xDB)
MID   = RGBColor(0x37, 0x41, 0x51)
LIGHT = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def f(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color: run.font.color.rgb = color

def border_bottom(p, color='1A56DB', sz='6'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '6'); bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)

def shd(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), fill)
    tcPr.append(s)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    f(r, 11, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    border_bottom(p, '0D1B2A', '6')

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    f(r, 10.5, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def item(text):
    p = doc.add_paragraph()
    m = p.add_run('✓  '); f(m, 10.5, bold=True, color=BLUE)
    t = p.add_run(text); f(t, 10.5, color=NAVY)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)

def sp(n=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(n)

# ━━━ ШАПКА ━━━
p = doc.add_paragraph()
r1 = p.add_run('Core'); f(r1, 24, bold=True, color=NAVY)
r2 = p.add_run('Nova'); f(r2, 24, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(1)
p2 = doc.add_paragraph()
r = p2.add_run('Цифровое ядро вашего бизнеса')
f(r, 9, color=LIGHT, italic=True)
p2.paragraph_format.space_after = Pt(14)
border_bottom(p2, '1A56DB', '2')
sp(10)

# ━━━ ЗАГОЛОВОК ━━━
p = doc.add_paragraph()
r = p.add_run('Стоимость разработки')
f(r, 22, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(3)

p = doc.add_paragraph()
r = p.add_run('Miss Miami 2026 — Цифровая платформа конкурса красоты')
f(r, 13, color=BLUE)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run(f'{datetime.date.today().strftime("%d.%m.%Y")}')
f(r, 9.5, color=LIGHT, italic=True)
p.paragraph_format.space_after = Pt(16)
sp(4)

# ━━━ ВВОДНЫЙ БЛОК ━━━
tbl0 = doc.add_table(rows=1, cols=1)
tbl0.style = 'Table Grid'
cell0 = tbl0.rows[0].cells[0]
cell0.width = Inches(6.4)
p = cell0.paragraphs[0]
r = p.add_run('Платформа Miss Miami 2026 охватывает весь цикл конкурса — от первой заявки участницы до финального вечера: регистрация, билеты, спонсоры, бренд и управление в одном месте.')
f(r, 10.5, color=NAVY)
p.paragraph_format.left_indent = Inches(0.1)
shd(cell0, 'EFF6FF')
sp(8)

# ━━━ ТАБЛИЦА ЦЕН ━━━
h1('Стоимость')
sp(4)

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
W = [Inches(1.4), Inches(3.6), Inches(1.2)]

for i, h in enumerate(['', 'Состав', 'Стоимость']):
    cell = tbl.rows[0].cells[i]; cell.width = W[i]
    p = cell.paragraphs[0]; r = p.add_run(h)
    f(r, 10, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd(cell, '0D1B2A')

rows_data = [
    ('Этап 1', 'Дизайн и все страницы сайта', '2 500 $'),
    ('Этап 2', 'Весь функционал платформы', '2 500 $'),
    ('ИТОГО', 'Полная платформа под ключ', '5 000 $'),
]
for ri, (phase, desc, price) in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    is_total = ri == 2
    fill = 'EFF6FF' if is_total else 'FFFFFF'
    for ci, text in enumerate([phase, desc, price]):
        cell = row.cells[ci]; cell.width = W[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        f(r, 10.5, bold=is_total, color=BLUE if is_total else NAVY)
        shd(cell, fill)
sp(10)

# ━━━ ЧТО ВХОДИТ ━━━
h1('Состав платформы')

h2('Для участниц')
for c in [
    'Онлайн-регистрация и подача заявки',
    'Загрузка фотографий и документов',
    'Оплата регистрационного взноса онлайн',
    'Личный кабинет со статусом заявки',
    '«Road to the Crown» — персональный путь подготовки из 8 шагов',
    'Автоматические письма: подтверждение, инструкции, напоминания, итоги',
]: item(c)

h2('Для зрителей')
for c in [
    'Покупка билетов онлайн',
    'Countdown-таймер до финала',
    'Подписка на новости Miss Miami',
]: item(c)

h2('Для спонсоров')
for c in [
    'Страница пакетов: Gold, Platinum, Титульный',
    'Sponsor Deck — PDF для скачивания',
    'Форма «Стать спонсором» — заявка или оплата онлайн',
]: item(c)

h2('О конкурсе')
for c in [
    'Правила, расписание, призы',
    'Благотворительная миссия Beauty with Purpose',
    'Галерея победительниц прошлых лет',
    'Профили судей',
    'Фото и видео с мероприятий',
    'Отзывы участниц',
]: item(c)

h2('Бренд и медиа')
for c in [
    'Медиа-кит и пресс-раздел',
    'Новости и блог (4 стартовые SEO-статьи: How to Prepare, Miss Miami Tips, Interview Preparation, Runway Training)',
    'SEO-оптимизация всех страниц',
    'Социальные сети',
]: item(c)

h2('Кабинет организатора')
for c in [
    'Все заявки в одном месте — статусы, фильтры, поиск',
    'Управление участницами: одобрить, отклонить, написать',
    'Email-рассылки и уведомления из кабинета',
    'Промокоды',
    'Голосование «People\'s Choice» с защитой от накрутки',
    'Магазин мерча Miss Miami',
    'Партнёрская программа (affiliate-ссылки)',
    'База клиентов — контакты автоматически попадают в систему',
]: item(c)

h2('Надёжность и доверие')
for c in [
    'FAQ, Контакты',
    'Политика конфиденциальности, Terms & Conditions, Refund Policy',
    'Адаптивный дизайн: десктоп, планшет, мобильный',
    '365 дней технического обслуживания после запуска',
]: item(c)

sp(8)

# ━━━ ОПЛАТА ━━━
h1('График платежей')
steps = [
    ('50%  —  2 500 $', 'предоплата до начала работ'),
    ('25%  —  1 250 $', 'после утверждения Этапа 1'),
    ('25%  —  1 250 $', 'после сдачи полной платформы'),
]
for amount, when in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(amount + '   '); f(r1, 11, bold=True, color=NAVY)
    r2 = p.add_run(when); f(r2, 10.5, color=LIGHT)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)

sp(12)

# ━━━ FOOTER ━━━
p = doc.add_paragraph()
border_bottom(p, '1A56DB', '2')
p.paragraph_format.space_after = Pt(6)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CoreNova  ·  info@corenova.ua  ·  lekki79.github.io/corenova')
f(r, 9, color=LIGHT, italic=True)

doc.save('/sessions/relaxed-jolly-pasteur/mnt/outputs/CoreNova_MissMiami_Стоимость_v2.docx')
print('Done!')
