from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

NAVY  = RGBColor(0x0D, 0x1B, 0x2A)
BLUE  = RGBColor(0x1A, 0x56, 0xDB)
MID   = RGBColor(0x37, 0x41, 0x51)
LIGHT = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x05, 0x96, 0x69)
RED   = RGBColor(0xDC, 0x26, 0x26)
LGRAY = RGBColor(0xF3, 0xF4, 0xF6)

def font(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_border_bottom(p, color_hex='1A56DB', sz='4'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), sz)
    bot.set(qn('w:space'), '6')
    bot.set(qn('w:color'), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

def shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    font(r, 13, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    add_border_bottom(p, '0D1B2A', '8')
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 11, bold=True, color=BLUE)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def body(text, color=None, italic=False, space_after=3):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, 10.5, color=color or MID, italic=italic)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, color=GREEN, mark='✓'):
    p = doc.add_paragraph()
    m = p.add_run(f'{mark}  ')
    font(m, 10.5, bold=True, color=color)
    t = p.add_run(text)
    font(t, 10.5, color=NAVY if mark == '✓' else LIGHT)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)

def chip(text):
    p = doc.add_paragraph()
    r = p.add_run(f'  {text}  ')
    font(r, 9, bold=True, color=BLUE)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)

def spacer(n=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(n)

# ━━━ HEADER ━━━
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r1 = p.add_run('Core')
font(r1, 26, bold=True, color=NAVY)
r2 = p.add_run('Nova')
font(r2, 26, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
r = p2.add_run('Цифровое ядро вашего бизнеса')
font(r, 10, color=LIGHT, italic=True)
p2.paragraph_format.space_after = Pt(16)

add_border_bottom(p2, '1A56DB', '2')
spacer(10)

# ━━━ TITLE ━━━
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('Коммерческое предложение')
font(r, 20, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run('Miss Miami 2026 — Полная цифровая платформа')
font(r, 13, color=BLUE)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
r = p.add_run(f'Дата: {datetime.date.today().strftime("%d.%m.%Y")}   |   Действует 30 дней')
font(r, 9.5, color=LIGHT, italic=True)
p.paragraph_format.space_after = Pt(20)

# ━━━ INTRO ━━━
h1('Суть предложения')
body('Miss Miami 2026 — это не просто сайт-визитка. Это цифровая платформа для проведения конкурса красоты: от регистрации участниц до продажи билетов, от управления заявками до автоматических email-уведомлений.', space_after=6)
body('CoreNova строит не набор отдельных страниц, а единую систему, где каждый компонент работает вместе и усиливает другой.', color=NAVY, space_after=6)
spacer()

# ━━━ ЦЕНА ━━━
h1('Инвестиция')

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
W = [Inches(2.6), Inches(2.8), Inches(1.2)]

for i, h in enumerate(['Этап', 'Состав', 'Стоимость']):
    cell = tbl.rows[0].cells[i]
    cell.width = W[i]
    p = cell.paragraphs[0]
    r = p.add_run(h)
    font(r, 10, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading(cell, '0D1B2A')

rows_data = [
    ('Этап 1', 'Дизайн и все страницы сайта\n(адаптив, GitHub Pages, анимации)', '2 500 $'),
    ('Этап 2', 'Функциональность и интеграции\n(auth, payments, CRM, email)', '2 500 $'),
    ('ИТОГО', '', '5 000 $'),
]
fills = ['FFFFFF', 'FFFFFF', 'EFF6FF']
for ri, (phase, desc, price) in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    is_total = ri == 2
    for ci, text in enumerate([phase, desc, price]):
        cell = row.cells[ci]
        cell.width = W[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        font(r, 10.5, bold=is_total, color=BLUE if is_total else NAVY)
        shading(cell, fills[ri])
spacer(8)

# ━━━ ЭТАП 1 ━━━
h1('Этап 1 — Дизайн и разработка сайта  |  2 500 $')
body('Срок: 3–4 недели  |  Статус: в работе', color=LIGHT, italic=True, space_after=8)

h2('Страницы и разделы')
pages = [
    'Главная — герой, миссия, countdown до финала, ценности, призыв к действию',
    'О нас — история, команда, миссия Beauty with Purpose',
    'Конкурс — правила, расписание, призы',
    'Участие — многошаговая форма с загрузкой фото/документов и оплатой',
    'Победительницы — галерея прошлых чемпионок',
    'Судьи — профили судей конкурса',
    'Галерея — фото и видео с мероприятий',
    'Спонсорам — пакеты Gold / Platinum / Титульный + PDF Sponsor Deck',
    'Билеты — интеграция с площадкой продажи билетов (Eventbrite или аналог)',
    'Медиа и пресс-раздел',
    'Новости / Блог (структура, 2 стартовые статьи)',
    'Отзывы участниц',
    'FAQ — вопросы и ответы по категориям',
    'Контакты + подписка на новости',
    'Политика конфиденциальности, Terms & Conditions, Refund Policy',
]
for pg in pages:
    bullet(pg)

h2('Технические характеристики')
specs = [
    'Полностью адаптивный дизайн: десктоп, планшет, мобильный',
    'Люксовая эстетика Miss Miami: золото, кремовый, бронза, Cinzel + Montserrat',
    'Плавные анимации, scroll-эффекты, быстрая загрузка (сжатые изображения)',
    'SEO-оптимизация страниц (мета-теги, заголовки, структура)',
    'Ссылки на социальные сети (Instagram, Facebook, TikTok)',
    'Хостинг GitHub Pages (бесплатно и стабильно)',
]
for sp in specs:
    bullet(sp)

# ━━━ ЭТАП 2 ━━━
h1('Этап 2 — Функциональность и интеграции  |  2 500 $')
body('Срок: 4–6 недель  |  Начало после утверждения Этапа 1', color=LIGHT, italic=True, space_after=8)

h2('Для участниц')
for c in [
    'Firebase: регистрация и авторизация (email + пароль)',
    'Форма заявки с загрузкой фотографий и документов (Firebase Storage)',
    'Stripe: оплата регистрационного взноса онлайн',
    'Личный кабинет участницы — статус заявки в реальном времени',
    '"Road to the Crown" — 8-шаговый путь подготовки участницы',
]:
    bullet(c)

h2('Для зрителей')
for c in [
    'Countdown-таймер до финала конкурса',
    'Подписка на новости (Mailchimp)',
    'Интеграция с площадкой продажи билетов',
]:
    bullet(c)

h2('Для спонсоров')
for c in [
    'Форма "Стать спонсором" — заявка или онлайн-оплата пакета',
    'PDF Sponsor Deck — готов для скачивания',
    'Stripe: оплата спонсорского пакета онлайн',
]:
    bullet(c)

h2('Автоматические email-цепочки (SendGrid)')
for c in [
    'Подтверждение регистрации участницы',
    'Напоминание о недостающих документах',
    'Уведомление о статусе заявки',
    'Инструкции и подготовка к мероприятию',
    'Благодарность после конкурса',
]:
    bullet(c)

h2('Кабинет организатора (Admin)')
for c in [
    'Все заявки участниц — просмотр, фильтрация, смена статуса',
    'Отправка email-уведомлений участницам из кабинета',
    'Промокоды — создание и управление (Stripe Coupons)',
    'Базовая аналитика: регистрации, платежи, статусы',
]:
    bullet(c)

h2('Голосование People\'s Choice')
for c in [
    'Простое голосование за участниц (Firebase)',
    'Защита от накрутки (по IP и email)',
    'Отображение результатов в реальном времени',
]:
    bullet(c)

# ━━━ ВНЕ РАМОК ━━━
h1('Не входит в данный договор — Будущие этапы')
body('Следующие функции выходят за рамки текущего бюджета и будут оцениваться отдельно:', space_after=6)
for o in [
    'Магазин мерча Miss Miami (отдельная платформа)',
    'Партнёрская/аффилейт-программа',
    'Внешняя CRM-интеграция (HubSpot, Salesforce)',
    'Написание SEO-статей для блога',
    'Дизайн медиа-кита (брендированный PDF)',
    'Мобильное приложение',
]:
    bullet(o, color=LIGHT, mark='—')

# ━━━ ЧТО НУЖНО ОТ КЛИЕНТА ━━━
h1('Что нам нужно от вас')
body('Для старта и работы над проектом нам необходимо:', space_after=6)
needed = [
    ('Логотип Miss Miami (SVG или PNG на прозрачном фоне)', 'Этап 1'),
    ('Фотографии команды (высокое разрешение)', 'Этап 1'),
    ('Фото с прошлых мероприятий для галереи', 'Этап 1'),
    ('Фотографии и имена прошлых победительниц', 'Этап 1'),
    ('Профили и фото судей', 'Этап 1'),
    ('Правила конкурса, расписание, информация о призах', 'Этап 1'),
    ('Sponsor Deck (или материалы для его создания)', 'Этап 1'),
    ('Доступ к аккаунту Stripe (создаёт Клиент)', 'Этап 2'),
    ('Доступ к аккаунту Firebase (создаёт Клиент)', 'Этап 2'),
    ('Выбор площадки для продажи билетов (Eventbrite и др.)', 'Этап 2'),
]

tbl2 = doc.add_table(rows=len(needed)+1, cols=2)
tbl2.style = 'Table Grid'
W2 = [Inches(4.2), Inches(1.4)]

for i, h in enumerate(['Что предоставить', 'Когда нужно']):
    cell = tbl2.rows[0].cells[i]
    cell.width = W2[i]
    p = cell.paragraphs[0]
    r = p.add_run(h)
    font(r, 10, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading(cell, '1A56DB')

for ri, (item, when) in enumerate(needed):
    row = tbl2.rows[ri+1]
    for ci, text in enumerate([item, when]):
        cell = row.cells[ci]
        cell.width = W2[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        font(r, 10, color=NAVY if ci == 0 else BLUE, bold=(ci==1))
        if ri % 2 == 0:
            shading(cell, 'F9FAFB')

spacer(8)

# ━━━ ОПЛАТА ━━━
h1('График платежей')
payments = [
    '50% — 2 500 $ — предоплата до начала работ',
    '25% — 1 250 $ — после завершения и приёмки Этапа 1',
    '25% — 1 250 $ — после завершения и приёмки Этапа 2',
]
for pay in payments:
    bullet(pay, color=BLUE, mark='→')

spacer(4)
body('Оплата: банковский перевод или PayPal. Реквизиты предоставляются отдельно.', color=LIGHT, italic=True)

# ━━━ ПРОЦЕСС ━━━
h1('Как мы работаем')
steps = [
    ('01  Старт', 'Получаем предоплату и материалы от Клиента. Начинаем работу.'),
    ('02  Разработка', 'Регулярные обновления и промежуточные показы. Вы видите прогрес на каждом этапе.'),
    ('03  Правки', 'До 3 раундов правок включены в стоимость каждого этапа.'),
    ('04  Приёмка', 'Клиент принимает этап, оплачивает следующую часть. Переходим к Этапу 2.'),
    ('05  Поддержка', 'После запуска — 30 дней бесплатной поддержки по техническим вопросам.'),
]
for step, desc in steps:
    p = doc.add_paragraph()
    r1 = p.add_run(step + '  ')
    font(r1, 10.5, bold=True, color=BLUE)
    r2 = p.add_run(desc)
    font(r2, 10.5, color=MID)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)

# ━━━ УСЛОВИЯ ━━━
h1('Условия договора')
terms = [
    'Исходный код и все файлы передаются Клиенту после финальной оплаты.',
    'Дополнительные правки сверх включённых — по договорённости.',
    'Срок начинается после получения предоплаты и необходимых материалов.',
    'Аккаунты Stripe, Firebase, Mailchimp — собственность Клиента.',
    'CoreNova не несёт ответственности за сторонние сервисы (Stripe, Firebase и др.).',
    'NDA подписывается по запросу Клиента.',
]
for t in terms:
    bullet(t, color=MID, mark='·')

# ━━━ ПОДПИСИ ━━━
spacer(10)
h1('Подписание')
body('Подписывая данный документ, стороны подтверждают согласие с условиями и объёмом работ.', space_after=16)

sig = doc.add_table(rows=4, cols=2)
sig.style = 'Table Grid'
data = [
    ['Клиент', 'Исполнитель'],
    ['Ирина Коваленко\nMiss Miami 2026', 'CoreNova\ninfo@corenova.ua'],
    ['Подпись: _________________________', 'Подпись: _________________________'],
    [f'Дата: ____________________________', 'Дата: ____________________________'],
]
fills2 = ['0D1B2A', '0D1B2A', 'FFFFFF', 'FFFFFF']
for ri, row_d in enumerate(data):
    for ci, text in enumerate(row_d):
        cell = sig.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        r = p.add_run(text)
        is_header = ri == 0
        font(r, 10.5, bold=is_header, color=WHITE if is_header else (NAVY if ri == 1 else LIGHT))
        shading(cell, fills2[ri])
        cell.width = Inches(2.8)

# ━━━ FOOTER ━━━
spacer(16)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CoreNova  |  info@corenova.ua  |  corenova.ua')
font(r, 9, color=LIGHT, italic=True)
add_border_bottom(p, '1A56DB', '2')

doc.save('/sessions/relaxed-jolly-pasteur/mnt/outputs/CoreNova_MissMiami2026_Предложение.docx')
print('Done!')
