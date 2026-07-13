from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

GOLD = RGBColor(0xC9, 0x92, 0x2A)
DARK = RGBColor(0x1A, 0x08, 0x00)
MID  = RGBColor(0x5A, 0x3A, 0x10)
LIGHT = RGBColor(0x8C, 0x60, 0x20)

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Georgia'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, 20, bold=True, color=DARK)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        # gold underline border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'C9922A')
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        set_font(run, 13, bold=True, color=GOLD)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    return p

def add_body(text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 11, color=MID)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p

def add_bullet(text, checkmark=True):
    p = doc.add_paragraph()
    mark = '✓  ' if checkmark else '—  '
    run = p.add_run(mark)
    set_font(run, 11, bold=True, color=GOLD)
    run2 = p.add_run(text)
    set_font(run2, 11, color=DARK)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_out_of_scope(text):
    p = doc.add_paragraph()
    run = p.add_run('✗  ')
    set_font(run, 11, bold=True, color=RGBColor(0x99,0x44,0x11))
    run2 = p.add_run(text)
    set_font(run2, 11, color=RGBColor(0x66,0x33,0x00))
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

# ── HEADER ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MISS MIAMI 2026')
set_font(run, 28, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Scope of Work & Service Agreement')
set_font(run2, 14, color=DARK)
p2.paragraph_format.space_after = Pt(2)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('─' * 52)
set_font(run3, 10, color=GOLD)
p3.paragraph_format.space_after = Pt(16)

# ── PARTIES ──
add_heading('Project Overview')
add_body('Client: Iryna Kovalenko, Miss Miami 2026')
add_body('Developer: Hella')
add_body(f'Date: {datetime.date.today().strftime("%B %d, %Y")}')
add_body('Website: lekki79.github.io/miss-miami-2026')

# ── PRICING TABLE ──
add_spacer()
add_heading('Investment Summary')

table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
widths = [Inches(2.8), Inches(2.4), Inches(1.4)]

headers = ['Phase', 'Description', 'Price']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.width = widths[i]
    p = cell.paragraphs[0]
    run = p.add_run(h)
    set_font(run, 11, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # gold background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'C9922A')
    tcPr.append(shd)

rows_data = [
    ('Phase 1', 'Website Design & Development\n(9 pages, responsive, GitHub Pages)', '$2,500'),
    ('Phase 2', 'Core Functionality\n(Auth, payments, personal account, admin)', '$2,500'),
    ('TOTAL', '', '$5,000'),
]
for r_idx, (phase, desc, price) in enumerate(rows_data):
    row = table.rows[r_idx + 1]
    is_total = r_idx == 2
    
    for c_idx, text in enumerate([phase, desc, price]):
        cell = row.cells[c_idx]
        cell.width = widths[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, 11, bold=is_total, color=GOLD if is_total else DARK)
        if is_total:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'FDF5E0')
            tcPr.append(shd)

add_spacer()

# ── PHASE 1 ──
add_heading('Phase 1 — Website Design & Development')
p = doc.add_paragraph()
run = p.add_run('Price: $2,500   |   Timeline: 3–4 weeks   |   Status: In Progress')
set_font(run, 10, italic=True, color=LIGHT)
p.paragraph_format.space_after = Pt(8)

add_heading('Included Pages', level=2)
pages = [
    'Home (index.html) — hero, mission teaser, values, CTA',
    'About — mission, team, timeline, values',
    'Apply — multi-step application form with payment placeholder',
    'Competition — rules, schedule, prizes',
    'Gallery — photo & video grid',
    'Sponsors — partner tiers and logos',
    'FAQ — categorized Q&A with search',
    'Contact — form and social links',
    'Winners — past champions showcase',
    'Judges — judge profiles',
    'Privacy Policy, Terms & Conditions, Refund Policy',
]
for p_text in pages:
    add_bullet(p_text)

add_heading('Technical Specifications', level=2)
specs = [
    'Fully responsive design (desktop, tablet, mobile)',
    'Luxury brand aesthetic — warm gold, cream, bronze palette',
    'Cinzel + Montserrat typography',
    'Smooth animations and scroll effects',
    'Optimized images for fast loading',
    'Hosted on GitHub Pages',
    'Cross-browser compatible',
]
for s in specs:
    add_bullet(s)

# ── PHASE 2 ──
add_heading('Phase 2 — Core Functionality')
p = doc.add_paragraph()
run = p.add_run('Price: $2,500   |   Timeline: 4–6 weeks   |   Starts after Phase 1 approval')
set_font(run, 10, italic=True, color=LIGHT)
p.paragraph_format.space_after = Pt(8)

add_heading('Contestant System', level=2)
contestant = [
    'Firebase authentication (email + password)',
    'Registration form with photo & video upload',
    'Stripe payment integration (registration fee)',
    'Personal dashboard — application status tracking',
    '"Road to the Crown" — 8-step contestant journey',
    'Email confirmation on registration (SendGrid)',
]
for c in contestant:
    add_bullet(c)

add_heading('Admin Panel', level=2)
admin = [
    'View and manage all applications',
    'Update contestant status (Pending / Approved / Rejected)',
    'Send email notifications to contestants',
    'Basic registration and payment reporting',
]
for a in admin:
    add_bullet(a)

add_heading('Additional Features', level=2)
extras = [
    'Email newsletter signup (Mailchimp integration)',
    'Sponsor/partner inquiry form with auto-response',
    'Contact form with email delivery',
]
for e in extras:
    add_bullet(e)

# ── OUT OF SCOPE ──
add_heading('Out of Scope (Future Phases)')
p = doc.add_paragraph()
run = p.add_run('The following features are NOT included in the current $5,000 agreement and will be quoted separately:')
set_font(run, 11, color=MID)
p.paragraph_format.space_after = Pt(8)

out_of_scope = [
    'Online voting system',
    'Ticket sales with QR code e-tickets',
    'Merch store',
    'Promo codes and discount management',
    'Full CRM with automated email sequences',
    'Blog / News section with CMS',
    'Advanced analytics dashboard',
    'Mobile app',
]
for o in out_of_scope:
    add_out_of_scope(o)

# ── PAYMENT SCHEDULE ──
add_heading('Payment Schedule')
payments = [
    '50% deposit ($2,500) — due before work begins',
    '25% ($1,250) — due upon Phase 1 completion',
    '25% ($1,250) — due upon Phase 2 completion',
]
for pay in payments:
    add_bullet(pay)

# ── TERMS ──
add_heading('Terms & Conditions')
terms = [
    'All content (text, photos, logos) to be provided by the Client.',
    'Stripe and Firebase accounts to be created and owned by the Client.',
    'Revisions: up to 3 rounds of revisions per phase included.',
    'Additional revisions billed at $75/hour.',
    'Timeline starts upon receipt of deposit and required materials.',
    'Source code and all files transferred to Client upon final payment.',
]
for t in terms:
    add_bullet(t)

# ── SIGNATURES ──
add_spacer()
add_heading('Agreement')
p = doc.add_paragraph()
run = p.add_run('By proceeding with payment, both parties agree to the terms outlined in this document.')
set_font(run, 11, color=MID)
p.paragraph_format.space_after = Pt(20)

# Signature table
sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'

sig_data = [
    ['Client: Iryna Kovalenko', 'Developer: Hella'],
    ['Signature: _____________________', 'Signature: _____________________'],
    [f'Date: ___________________________', 'Date: ___________________________'],
]
for r_idx, row_data in enumerate(sig_data):
    for c_idx, text in enumerate(row_data):
        cell = sig_table.rows[r_idx].cells[c_idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        bold = r_idx == 0
        set_font(run, 11, bold=bold, color=DARK if bold else MID)
        cell.width = Inches(2.8)

doc.save('/sessions/relaxed-jolly-pasteur/mnt/outputs/Miss_Miami_2026_Scope_of_Work.docx')
print("Done!")
