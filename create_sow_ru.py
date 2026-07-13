from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.2)
section.right_margin = Inches(1.2)

GOLD  = RGBColor(0xC9, 0x92, 0x2A)
DARK  = RGBColor(0x1A, 0x08, 0x00)
MID   = RGBColor(0x5A, 0x3A, 0x10)
LIGHT = RGBColor(0x8C, 0x60, 0x20)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
OUT_C = RGBColor(0x99, 0x44, 0x11)

def font(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Georgia'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        font(run, 18, bold=True, color=DARK)
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(5)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), 'C9922A')
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        font(run, 12, bold=True, color=GOLD)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    font(run, 11, color=MID, italic=italic)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, ok=True):
    p = doc.add_paragraph()
    mark = p.add_run(('+ ' if ok else 'x '))
    font(mark, 11, bold=True, color=GOLD if ok else OUT_C)
    txt = p.add_run(text)
    font(txt, 11, color=DARK if ok else OUT_C)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)

def shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

# ── ШАПКА ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MISS MIAMI 2026')
font(run, 28, bold=True, color=GOLD)
p.paragraph_format.space_after = Pt(2)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Техническое задание и договор об оказании услуг')
font(run2, 13, color=DARK)
p2.paragraph_format.space_after = Pt(2)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('=' * 55)
font(run3, 9, color=GOLD)
p3.paragraph_format.space_after = Pt(14)

# ── СТОРОНЫ ──
heading('Общая информация')
body('Клиент: Ирина Коваленко, Miss Miami 2026')
body('Исполнитель: CoreNova')
body(f'Дата: {datetime.date.today().strftime("%d.%m.%Y")}')
body('Сайт: lekki79.github.io/miss-miami-2026')

# ── ТАБЛИЦА ЦЕН ──
doc.add_paragraph()
heading('Стоимость работ')

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
W = [Inches(2.8), Inches(2.4), Inches(1.4)]

for i, h in enumerate(['Этап', 'Описание', 'Стоимость']):
    cell = tbl.rows[0].cells[i]
    cell.width = W[i]
    p = cell.paragraphs[0]
    r = p.add_run(h)
    font(r, 11, bold=True, color=WHITE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading(cell, 'C9922A')

rows_data = [
    ('Этап 1', 'Дизайн и разработка сайта\n(9 страниц, адаптив, GitHub Pages)', '2 500 $'),
    ('Этап 2', 'Функциональность\n(Авторизация, оплата, кабинет, админка)', '2 500 $'),
    ('ИТОГО', '', '5 000 $'),
]
for ri, (phase, desc, price) in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    is_total = ri == 2
    for ci, text in enumerate([phase, desc, price]):
        cell = row.cells[ci]
        cell.width = W[ci]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        font(r, 11, bold=is_total, color=GOLD if is_total else DARK)
        if is_total:
            shading(cell, 'FDF5E0')

doc.add_paragraph()

# ── ЭТАП 1 ──
heading('Этап 1 — Дизайн и разработка сайта')
body('Стоимость: 2 500 $   |   Срок: 3–4 недели   |   Статус: в работе', italic=True)

heading('Страницы сайта', level=2)
for pg in [
    'Главная — герой, миссия, ценности, призыв к действию',
    'О нас — миссия, команда, история, ценности',
    'Участие — многошаговая форма заявки с оплатой',
    'Конкурс — правила, расписание, призы',
    'Галерея — фото и видео',
    'Спонсоры — уровни партнёрства и логотипы',
    'FAQ — вопросы и ответы с поиском',
    'Контакты — форма и ссылки на соцсети',
    'Победительницы — архив чемпионок',
    'Судьи — профили судей',
    'Политика конфиденциальности, Условия использования, Политика возврата',
]:
    bullet(pg)

heading('Технические требования', level=2)
for sp in [
    'Полностью адаптивный дизайн (десктоп, планшет, мобильный)',
    'Люксовая эстетика — золото, кремовый, бронза',
    'Шрифты Cinzel + Montserrat',
    'Анимации и плавные переходы',
    'Оптимизированные изображения для быстрой загрузки',
    'Хостинг на GitHub Pages',
    'Совместимость с основными браузерами',
]:
    bullet(sp)

# ── ЭТАП 2 ──
heading('Этап 2 — Функциональность')
body('Стоимость: 2 500 $   |   Срок: 4–6 недель   |   Начало после утверждения Этапа 1', italic=True)

heading('Система регистрации участниц', level=2)
for c in [
    'Авторизация через Firebase (email + пароль)',
    'Форма регистрации с загрузкой фото и видео',
    'Интеграция оплаты через Stripe (взнос за участие)',
    'Личный кабинет — статус рассмотрения заявки',
    '"Road to the Crown" — 8 шагов подготовки участницы',
    'Email-подтверждение при регистрации (SendGrid)',
]:
    bullet(c)

heading('Административная панель', level=2)
for a in [
    'Просмотр и управление всеми заявками',
    'Изменение статуса заявки (На рассмотрении / Принята / Отклонена)',
    'Отправка email-уведомлений участницам',
    'Базовые отчёты по регистрациям и платежам',
]:
    bullet(a)

heading('Дополнительный функционал', level=2)
for e in [
    'Подписка на email-рассылку (интеграция с Mailchimp)',
    'Форма для спонсоров и партнёров с автоответом',
    'Форма обратной связи с доставкой на email',
]:
    bullet(e)

# ── ВНЕ РАМОК ──
heading('Не входит в данный договор (Будущие этапы)')
body('Следующие функции НЕ включены в текущую сумму 5 000 $ и будут оцениваться отдельно:')
for o in [
    'Онлайн-голосование',
    'Продажа билетов с QR-кодами',
    'Магазин мерча',
    'Промокоды и скидки',
    'Полная CRM с автоматическими цепочками писем',
    'Раздел новостей / блог с системой управления контентом',
    'Расширенная аналитика',
    'Мобильное приложение',
]:
    bullet(o, ok=False)

# ── ОПЛАТА ──
heading('График платежей')
for pay in [
    '50% предоплата (2 500 $) — до начала работ',
    '25% (1 250 $) — после завершения Этапа 1',
    '25% (1 250 $) — после завершения Этапа 2',
]:
    bullet(pay)

# ── УСЛОВИЯ ──
heading('Условия сотрудничества')
for t in [
    'Весь контент (тексты, фото, логотипы) предоставляет Клиент.',
    'Аккаунты Stripe и Firebase создаются и остаются собственностью Клиента.',
    'Правки: до 3 раундов правок на каждый этап включены в стоимость.',
    'Дополнительные правки оплачиваются отдельно по договорённости.',
    'Срок начинается после получения предоплаты и необходимых материалов.',
    'Исходный код и все файлы передаются Клиенту после финальной оплаты.',
]:
    bullet(t)

# ── ПОДПИСИ ──
doc.add_paragraph()
heading('Подписание договора')
body('Подписывая данный документ, стороны подтверждают согласие с указанными условиями.')
doc.add_paragraph()

sig = doc.add_table(rows=3, cols=2)
sig.style = 'Table Grid'
data = [
    ['Клиент: Ирина Коваленко', 'Исполнитель: CoreNova'],
    ['Подпись: _______________________', 'Подпись: _______________________'],
    ['Дата: ___________________________', 'Дата: ___________________________'],
]
for ri, row_d in enumerate(data):
    for ci, text in enumerate(row_d):
        cell = sig.rows[ri].cells[ci]
        p = cell.paragraphs[0]
        r = p.add_run(text)
        font(r, 11, bold=(ri == 0), color=DARK if ri == 0 else MID)
        cell.width = Inches(2.8)

doc.save('/sessions/relaxed-jolly-pasteur/mnt/outputs/Miss_Miami_2026_Договор.docx')
print('Done!')
